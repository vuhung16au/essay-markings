"""Streamlit frontend for the PTE essay marker."""

from __future__ import annotations

import json
import os
from io import BytesIO
from textwrap import dedent
from pathlib import Path
from urllib import error, request

import streamlit as st


ROOT_DIR = Path(__file__).resolve().parents[1]

CRITERIA_LABELS = {
    "content": "Content",
    "development_structure_coherence": "Development, Structure & Coherence",
    "form": "Form",
    "grammar": "Grammar",
    "linguistic_range": "Linguistic Range",
    "spelling": "Spelling",
    "vocabulary": "Vocabulary",
}


def load_styles(theme_mode: str) -> None:
    css_path = Path(__file__).parent / "assets" / "styles.css"
    if css_path.exists():
        theme_variables = dedent(
            """
            <style>
            :root {
              --bg: #0d1117;
              --bg-accent: #1a2230;
              --paper: rgba(16, 22, 31, 0.88);
              --paper-strong: #131a24;
              --ink: #f3efe7;
              --muted: #b8c0cc;
              --accent: #ff8a4c;
              --accent-soft: rgba(255, 138, 76, 0.18);
              --line: rgba(255, 255, 255, 0.12);
              --shadow: 0 24px 70px rgba(0, 0, 0, 0.35);
            }
            </style>
            """
        ).strip()

        if theme_mode == "Light":
            theme_variables = dedent(
                """
                <style>
                :root {
                  --bg: #f4ede1;
                  --bg-accent: #fff8ef;
                  --paper: rgba(255, 252, 246, 0.9);
                  --paper-strong: #fffaf2;
                  --ink: #1f2933;
                  --muted: #5d6a75;
                  --accent: #b85c38;
                  --accent-soft: rgba(184, 92, 56, 0.16);
                  --line: rgba(103, 80, 64, 0.18);
                  --shadow: 0 24px 70px rgba(118, 82, 45, 0.14);
                }
                </style>
                """
            ).strip()

        stylesheet = f"{theme_variables}\n<style>\n{css_path.read_text().strip()}\n</style>"
        st.markdown(stylesheet, unsafe_allow_html=True)


def load_local_env() -> dict[str, str]:
    values: dict[str, str] = {}
    env_path = ROOT_DIR / ".env.local"
    if not env_path.exists():
        return values

    for line in env_path.read_text().splitlines():
        cleaned = line.strip()
        if not cleaned or cleaned.startswith("#") or "=" not in cleaned:
            continue
        key, value = cleaned.split("=", 1)
        values[key.strip()] = value.strip()
    return values


@st.cache_data(show_spinner=False)
def cached_sample_data() -> tuple[list[dict], list[dict]]:
    return load_sample_data()


def get_backend_url() -> str:
    env_values = load_local_env()
    host = os.getenv("BACKEND_HOST") or env_values.get("BACKEND_HOST", "0.0.0.0")
    port = os.getenv("BACKEND_PORT") or env_values.get("BACKEND_PORT", "8000")
    if host == "0.0.0.0":
        host = "localhost"
    return f"http://{host}:{port}"


def load_sample_data() -> tuple[list[dict], list[dict]]:
    questions_path = ROOT_DIR / "data" / "sample_questions.json"
    essays_path = ROOT_DIR / "data" / "sample_essays.json"
    questions = json.loads(questions_path.read_text()) if questions_path.exists() else []
    essays = json.loads(essays_path.read_text()) if essays_path.exists() else []
    return questions, essays


def build_sample_options() -> dict[str, tuple[str, str]]:
    questions, essays = cached_sample_data()
    question_lookup = {item["id"]: item["question"] for item in questions}
    options: dict[str, tuple[str, str]] = {}
    for essay in essays:
        label = f"Sample {essay['id']} - {essay['quality_level'].replace('_', ' ').title()}"
        options[label] = (question_lookup.get(essay["question_id"], ""), essay["text"])
    return options


def count_words(text: str) -> int:
    return len(text.split())


def submit_for_grading(question: str, essay: str) -> dict:
    payload = json.dumps({"question": question, "essay": essay}).encode("utf-8")
    endpoint = f"{get_backend_url()}/api/grade-essay"
    http_request = request.Request(
        endpoint,
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with request.urlopen(http_request, timeout=90) as response:
            return json.loads(response.read().decode("utf-8"))
    except error.HTTPError as exc:
        try:
            body = json.loads(exc.read().decode("utf-8"))
            detail = body.get("detail", str(exc))
        except json.JSONDecodeError:
            detail = str(exc)
        raise RuntimeError(f"Backend request failed: {detail}") from exc
    except error.URLError as exc:
        raise RuntimeError(
            f"Could not reach the backend at {get_backend_url()}. Start the FastAPI server and try again."
        ) from exc


def backend_is_healthy() -> bool:
    endpoint = f"{get_backend_url()}/health"
    try:
        with request.urlopen(endpoint, timeout=5) as response:
            payload = json.loads(response.read().decode("utf-8"))
            return payload.get("status") == "ok"
    except Exception:
        return False


def render_scores(scores: dict) -> None:
    rows = []
    total = 0
    total_max = 0
    for key, label in CRITERIA_LABELS.items():
        item = scores[key]
        total += item["score"]
        total_max += item["max"]
        rows.append({"Criterion": label, "Score": f"{item['score']} / {item['max']}"})

    st.markdown(f"**Overall score:** `{total} / {total_max}`")
    st.table(rows)


def render_feedback(feedback: dict) -> None:
    rows = [{"Area": area.title(), "Details": details} for area, details in feedback.items()]
    st.table(rows)


def render_points(title: str, points: list[str]) -> None:
    st.markdown(f"**{title}**")
    for point in points:
        st.write(f"- {point}")


def render_details(details: dict, scores: dict) -> None:
    st.markdown("**Detail**")
    for key, label in CRITERIA_LABELS.items():
        item = details[key]
        score = scores[key]
        with st.expander(f"{label} ({score['score']} / {score['max']})", expanded=False):
            st.write(item["analysis"])
            deductions = item.get("deductions") or []
            if deductions:
                st.markdown("**Why marks were deducted**")
                for deduction in deductions:
                    st.write(f"- {deduction}")
            else:
                st.caption("No deductions were applied in this category.")


def total_score(scores: dict) -> tuple[int, int]:
    earned = sum(item["score"] for item in scores.values())
    maximum = sum(item["max"] for item in scores.values())
    return earned, maximum


def build_markdown_report(question: str, essay: str, result: dict) -> str:
    earned, maximum = total_score(result["scores"])
    lines = [
        "# Essay Report",
        "",
        "## Essay question",
        question.strip(),
        "",
        "## Submitted essay",
        essay.strip(),
        "",
        "## Overall score",
        f"{earned} / {maximum}",
        "",
        "## Detailed breakdown of scores",
        "",
        "| Category | Score |",
        "| --- | --- |",
    ]

    for key, label in CRITERIA_LABELS.items():
        score = result["scores"][key]
        lines.append(f"| {label} | {score['score']} / {score['max']} |")

    lines.extend(
        [
            "",
            "## Comments and feedback for each category",
            "",
        ]
    )

    details = result.get("details") or {}
    for key, label in CRITERIA_LABELS.items():
        score = result["scores"][key]
        lines.extend(
            [
                f"### {label} ({score['score']} / {score['max']})",
                (details.get(key) or {}).get("analysis", "No detailed analysis available."),
            ]
        )

        deductions = (details.get(key) or {}).get("deductions") or []
        if deductions:
            lines.append("")
            lines.append("Reasons for deducted marks:")
            for deduction in deductions:
                lines.append(f"- {deduction}")

        if key in result["feedback"]:
            lines.append("")
            lines.append(f"Feedback note: {result['feedback'][key]}")

        lines.append("")

    lines.extend(
        [
            "## Suggestions for improvement",
            "",
        ]
    )
    for item in result.get("improvements", []):
        lines.append(f"- {item}")

    good_points = result.get("good_points") or []
    if good_points:
        lines.extend(
            [
                "",
                "## Strengths",
                "",
            ]
        )
        for item in good_points:
            lines.append(f"- {item}")

    return "\n".join(lines).strip() + "\n"


def build_docx_report(question: str, essay: str, result: dict) -> bytes:
    from docx import Document

    earned, maximum = total_score(result["scores"])
    document = Document()
    document.add_heading("Essay Report", level=0)

    document.add_heading("Essay question", level=1)
    document.add_paragraph(question.strip())

    document.add_heading("Submitted essay", level=1)
    document.add_paragraph(essay.strip())

    document.add_heading("Overall score", level=1)
    document.add_paragraph(f"{earned} / {maximum}")

    document.add_heading("Detailed breakdown of scores", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Score"
    for key, label in CRITERIA_LABELS.items():
        row_cells = table.add_row().cells
        row_cells[0].text = label
        score = result["scores"][key]
        row_cells[1].text = f"{score['score']} / {score['max']}"

    document.add_heading("Comments and feedback for each category", level=1)
    details = result.get("details") or {}
    for key, label in CRITERIA_LABELS.items():
        score = result["scores"][key]
        section = details.get(key) or {}
        document.add_heading(f"{label} ({score['score']} / {score['max']})", level=2)
        document.add_paragraph(section.get("analysis", "No detailed analysis available."))

        deductions = section.get("deductions") or []
        if deductions:
            document.add_paragraph("Reasons for deducted marks:")
            for deduction in deductions:
                document.add_paragraph(deduction, style="List Bullet")

        if key in result["feedback"]:
            document.add_paragraph(f"Feedback note: {result['feedback'][key]}")

    document.add_heading("Suggestions for improvement", level=1)
    for item in result.get("improvements", []):
        document.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _wrap_text(text: str, width: int = 95) -> list[str]:
    paragraphs = text.splitlines() or [text]
    wrapped: list[str] = []
    for paragraph in paragraphs:
        stripped = paragraph.strip()
        if not stripped:
            wrapped.append("")
            continue
        wrapped.extend(dedent(stripped).splitlines())
    return wrapped


def build_pdf_report(question: str, essay: str, result: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.pdfgen import canvas

    def wrap_lines(text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
        output: list[str] = []
        for paragraph in text.splitlines():
            if not paragraph.strip():
                output.append("")
                continue
            words = paragraph.split()
            current = words[0]
            for word in words[1:]:
                trial = f"{current} {word}"
                if stringWidth(trial, font_name, font_size) <= max_width:
                    current = trial
                else:
                    output.append(current)
                    current = word
            output.append(current)
        return output

    earned, maximum = total_score(result["scores"])
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    left = 50
    top = height - 50
    line_height = 15
    max_width = width - 2 * left
    y = top

    def ensure_space(lines_needed: int = 1) -> None:
        nonlocal y
        if y - (lines_needed * line_height) < 50:
            pdf.showPage()
            y = top

    def draw_block(text: str, font_name: str = "Helvetica", font_size: int = 11) -> None:
        nonlocal y
        lines = wrap_lines(text, font_name, font_size, max_width)
        ensure_space(max(1, len(lines)))
        pdf.setFont(font_name, font_size)
        for line in lines:
            pdf.drawString(left, y, line)
            y -= line_height

    draw_block("Essay Report", "Helvetica-Bold", 16)
    y -= 5
    draw_block("Essay question", "Helvetica-Bold", 13)
    draw_block(question.strip())
    y -= 5
    draw_block("Submitted essay", "Helvetica-Bold", 13)
    draw_block(essay.strip())
    y -= 5
    draw_block("Overall score", "Helvetica-Bold", 13)
    draw_block(f"{earned} / {maximum}")
    y -= 5
    draw_block("Detailed breakdown of scores", "Helvetica-Bold", 13)
    for key, label in CRITERIA_LABELS.items():
        score = result["scores"][key]
        draw_block(f"{label}: {score['score']} / {score['max']}", "Helvetica", 11)

    y -= 5
    draw_block("Comments and feedback for each category", "Helvetica-Bold", 13)
    details = result.get("details") or {}
    for key, label in CRITERIA_LABELS.items():
        score = result["scores"][key]
        section = details.get(key) or {}
        draw_block(f"{label} ({score['score']} / {score['max']})", "Helvetica-Bold", 12)
        draw_block(section.get("analysis", "No detailed analysis available."))
        deductions = section.get("deductions") or []
        if deductions:
            draw_block("Reasons for deducted marks:", "Helvetica-Bold", 11)
            for deduction in deductions:
                draw_block(f"- {deduction}")
        if key in result["feedback"]:
            draw_block(f"Feedback note: {result['feedback'][key]}")
        y -= 3

    draw_block("Suggestions for improvement", "Helvetica-Bold", 13)
    for item in result.get("improvements", []):
        draw_block(f"- {item}")

    pdf.save()
    return buffer.getvalue()


def render_export_buttons(question: str, essay: str, result: dict) -> None:
    st.markdown("**Export Essay Report**")
    markdown_report = build_markdown_report(question, essay, result)
    markdown_col, word_col, pdf_col = st.columns(3)

    with markdown_col:
        st.download_button(
            "Export as Markdown",
            data=markdown_report.encode("utf-8"),
            file_name="essay-report.md",
            mime="text/markdown",
            use_container_width=True,
        )

    with word_col:
        try:
            word_report = build_docx_report(question, essay, result)
            st.download_button(
                "Export as Word",
                data=word_report,
                file_name="essay-report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.button("Export as Word", disabled=True, use_container_width=True)
            st.caption("Install `python-docx` to enable Word export.")

    with pdf_col:
        try:
            pdf_report = build_pdf_report(question, essay, result)
            st.download_button(
                "Export as PDF",
                data=pdf_report,
                file_name="essay-report.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except ImportError:
            st.button("Export as PDF", disabled=True, use_container_width=True)
            st.caption("Install `reportlab` to enable PDF export.")


st.set_page_config(page_title="PTE Essay Marker", layout="wide")
if "question" not in st.session_state:
    st.session_state.question = ""
if "essay" not in st.session_state:
    st.session_state.essay = ""
if "result" not in st.session_state:
    st.session_state.result = None
if "last_error" not in st.session_state:
    st.session_state.last_error = None
if "theme_mode" not in st.session_state:
    st.session_state.theme_mode = "Dark"

load_styles(st.session_state.theme_mode)

hero_col, toggle_col = st.columns([5, 1.5])
with hero_col:
    st.title("PTE Essay Marker")
    st.caption("Submit a PTE essay question and response to receive structured grading feedback from the backend.")
with toggle_col:
    st.markdown("**Theme**")
    st.radio(
        "Theme",
        ["Dark", "Light"],
        horizontal=True,
        key="theme_mode",
        label_visibility="collapsed",
    )

if backend_is_healthy():
    st.success("Backend connection is live.")
else:
    st.warning("Backend is not reachable right now. You can still draft an essay, but grading will fail until the API is running.")

sample_options = build_sample_options()
selected_sample = st.selectbox(
    "Load a sample response",
    ["Custom entry", *sample_options.keys()],
    index=0,
)

if selected_sample != "Custom entry":
    sample_question, sample_essay = sample_options[selected_sample]
    st.session_state.question = sample_question
    st.session_state.essay = sample_essay

question = st.text_area(
    "Essay question",
    height=120,
    key="question",
    placeholder="Paste the PTE prompt here...",
)
essay = st.text_area(
    "Essay response",
    height=320,
    key="essay",
    placeholder="Write or paste the essay here...",
)

word_count = count_words(essay)
st.caption(f"Essay word count: {word_count}")
if essay and word_count < 200:
    st.warning("PTE essays are typically stronger at 200+ words. You can still submit, but the response may score lower on form and development.")

if st.button("Submit for grading", type="primary"):
    if not question.strip() or not essay.strip():
        st.session_state.last_error = "Please provide both the essay question and essay text."
        st.error("Please provide both the essay question and essay text.")
    else:
        with st.spinner("Sending essay to the grading API..."):
            try:
                st.session_state.result = submit_for_grading(question.strip(), essay.strip())
                st.session_state.last_error = None
                st.success("Essay graded successfully.")
            except RuntimeError as exc:
                st.session_state.result = None
                st.session_state.last_error = str(exc)
                st.error(str(exc))

result = st.session_state.result
if result:
    st.subheader("Results")
    left_col, right_col = st.columns([1.1, 1])
    with left_col:
        st.markdown("**Score breakdown**")
        render_scores(result["scores"])
    with right_col:
        st.markdown("**Feedback**")
        render_feedback(result["feedback"])

    render_points("Good points", result["good_points"])
    render_points("Improvements", result["improvements"])
    if result.get("details"):
        render_details(result["details"], result["scores"])
    render_export_buttons(question, essay, result)
else:
    st.subheader("Results")
    message = "Submit an essay to see score breakdowns, targeted feedback, and improvement suggestions."
    if st.session_state.last_error:
        message = f"{message} The latest request failed, but you can edit and resubmit."
    st.info(message)
